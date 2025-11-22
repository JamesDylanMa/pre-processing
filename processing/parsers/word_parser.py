"""
Word document parser
"""
from typing import Dict
from pathlib import Path
try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


class WordParser:
    """Parse Word documents"""
    
    def __init__(self):
        self.name = "word_parser"
    
    def parse(self, file_path: str) -> Dict:
        """
        Parse Word file and extract content
        
        Args:
            file_path: Path to Word file
            
        Returns:
            Dictionary with extracted content
        """
        result = {
            "file_path": file_path,
            "parser": self.name,
            "paragraphs": [],
            "tables": [],
            "text": "",
            "metadata": {}
        }
        
        if not DOCX_AVAILABLE:
            result["error"] = "python-docx library not available"
            return result
        
        try:
            doc = Document(file_path)
            
            # Extract paragraphs
            paragraphs = []
            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs.append(para.text)
                    result["paragraphs"].append({
                        "text": para.text,
                        "style": para.style.name if para.style else None
                    })
            
            # Extract tables
            for i, table in enumerate(doc.tables):
                table_data = {
                    "table_number": i + 1,
                    "rows": []
                }
                for row in table.rows:
                    row_data = [cell.text for cell in row.cells]
                    table_data["rows"].append(row_data)
                result["tables"].append(table_data)
            
            result["text"] = "\n\n".join(paragraphs)
            result["metadata"]["total_paragraphs"] = len(paragraphs)
            result["metadata"]["total_tables"] = len(doc.tables)
            
            # Extract core properties if available
            if doc.core_properties:
                result["metadata"]["document_properties"] = {
                    "title": doc.core_properties.title,
                    "author": doc.core_properties.author,
                    "created": str(doc.core_properties.created) if doc.core_properties.created else None,
                    "modified": str(doc.core_properties.modified) if doc.core_properties.modified else None
                }
        
        except Exception as e:
            result["error"] = f"Failed to parse Word document: {str(e)}"
        
        return result


